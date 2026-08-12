# -*- coding: utf-8 -*-
"""
执法处建筑市场检查周报材料 自动生成程序
=========================================

输入数据（口径已经样例文档/百日攻坚xlsx交叉验证）：

1. 检查原始数据导出（平台导出的 HTML 格式 .xls，如 检查数据0803.xls）
   - 本周/上周 全市、各区、执法处 建筑市场检查单数/发现问题数/责改立案数：
     与“百日攻坚统计表生成程序(deploy)”同口径（检查单名称模板 + 添加时间
     + 机构名称子串划区），已对 7.20-7.26 / 7.27-8.2 两周逐行验证与
     百日攻坚xlsx 完全一致，因此不再需要百日攻坚xlsx 作为输入
   - 市级“立案数”与前八区“拟立案区”：任务名称含“建筑市场”、机构->区/市级
     按基本口径映射、处理结果为 一般行政处罚/简易行政处罚
2. 处罚原始数据导出（如 处罚数据0803.xls）
   - 本周处罚总数：一般+简易处罚决定书，按添加日期落在本周
   - 每百工程处罚量排名：前八区 处罚件数/在施工程量 降序
3. 2025年建筑市场检查.xls（固定）：去年同期检查数/问题数（按检查开始日期）
4. 2025年工程建设领域处罚.xls（固定）：去年同期处罚件数（按决定日期）
5. 基本口径.xls（固定）：在施工程量、机构->行政区映射

上周（环比）数值优先取“周报历史记录.json”中上周生成时保存的值（与上周报告一致），
没有历史记录时用原始导出重算（平台补录可能略有出入）。

运行：双击 start_gui.bat，或用 Python312 运行本脚本。
"""

import datetime as dt
import glob
import json
import os
import sys
import traceback
import warnings

import pandas as pd

warnings.filterwarnings("ignore")

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# 数据目录（历史记录）：默认与程序同目录；Linux/Docker 部署时可设 REPORT_DATA_DIR 持久化
DATA_DIR = os.environ.get("REPORT_DATA_DIR", SCRIPT_DIR)
HISTORY_PATH = os.path.join(DATA_DIR, "周报历史记录.json")

FIXED_CHECK_XLS = os.path.join(SCRIPT_DIR, "2025年建筑市场检查.xls")
FIXED_PUNISH_XLS = os.path.join(SCRIPT_DIR, "2025年工程建设领域处罚.xls")
FIXED_KOUJING_XLS = os.path.join(SCRIPT_DIR, "基本口径.xls")

# 处理结果口径
RESULT_PROBLEM = {"其他", "责令改正", "一般行政处罚", "简易行政处罚"}   # 发现问题
RESULT_ZGL = {"责令改正", "一般行政处罚", "简易行政处罚"}               # 责改和立案
RESULT_LIAN = {"一般行政处罚", "简易行政处罚"}                          # 立案
PUNISH_DOCS = {"一般处罚决定书", "简易处罚决定书"}                      # 处罚件数

DIST_SHORT = {"经济技术开发区": "经开区"}

# 建筑市场检查单模板（与 deploy/百日攻坚xlsx 口径一致）
BM_TEMPLATES = {"房屋市政工程建筑市场及关键人员履职检查单（现场检查）",
                "房屋市政工程建筑市场行为检查",
                "房屋市政工程建筑市场行为检查（非现场检查）"}

# 机构名称子串 -> 区（deploy 口径；均未命中的视为 执法处/市级）
DIST_SUB_MAP = {"东城": "东城区", "西城": "西城区", "朝阳": "朝阳区", "海淀": "海淀区",
                "丰台": "丰台区", "石景山": "石景山区", "门头沟": "门头沟区", "房山": "房山区",
                "大兴": "大兴区", "通州": "通州区", "顺义": "顺义区", "平谷区": "平谷区",
                "昌平": "昌平区", "怀柔": "怀柔区", "密云": "密云区", "延庆": "延庆区",
                "经济技术开发区": "经济技术开发区", "经开区": "经济技术开发区"}

# 区行序（同百日攻坚xlsx，用于拟立案区并列时的排序）
DISTRICTS_ORDER = ["东城区", "西城区", "朝阳区", "海淀区", "丰台区", "石景山区",
                   "门头沟区", "房山区", "大兴区", "通州区", "顺义区", "平谷区",
                   "昌平区", "怀柔区", "密云区", "延庆区", "经济技术开发区"]


# ---------------------------------------------------------------------------
# 基础工具
# ---------------------------------------------------------------------------
_READ_CACHE = {}


def read_any_excel(path, **kw):
    """读取平台导出文件：自动识别 HTML 伪装的 .xls 与真正的 xls/xlsx。

    按 (路径, 修改时间) 缓存：一次生成中多次读取同一导出文件时只解析一遍。
    返回副本，调用方可安全修改。
    """
    key = (os.path.abspath(path), os.path.getmtime(path))
    if key not in _READ_CACHE:
        with open(path, "rb") as f:
            head = f.read(8)
        if head.startswith(b"\xef\xbb\xbf") or head.lstrip().startswith(b"<"):
            df = pd.read_html(path, encoding="utf-8")[0]
        else:
            df = pd.read_excel(path, **kw)
        if len(_READ_CACHE) >= 8:
            _READ_CACHE.pop(next(iter(_READ_CACHE)))
        _READ_CACHE[key] = df
    return _READ_CACHE[key].copy()


def extract_district_sub(org):
    """deploy 口径：机构名称子串 -> 区名；均未命中的视为 执法处（市级）。"""
    s = str(org)
    for k, v in DIST_SUB_MAP.items():
        if k in s:
            return v
    return "执法处"


def contains_any(series, keys):
    """处理结果包含任一关键词（子串匹配，与 deploy 口径一致）。"""
    s = series.astype(str)
    m = pd.Series(False, index=series.index)
    for k in keys:
        m = m | s.str.contains(k, na=False, regex=False)
    return m


def md_mask(series, m0, d0, m1, d1):
    """按 月-日 区间过滤（支持跨年周），series 为日期列。"""
    s = pd.to_datetime(series, errors="coerce")
    md = s.dt.month * 100 + s.dt.day
    lo, hi = m0 * 100 + d0, m1 * 100 + d1
    if lo <= hi:
        return (md >= lo) & (md <= hi)
    return (md >= lo) | (md <= hi)


def day_mask(series, d0, d1):
    """按完整日期区间过滤（含当日全天）。"""
    s = pd.to_datetime(series, errors="coerce")
    t0, t1 = pd.Timestamp(d0), pd.Timestamp(d1)
    return (s >= t0) & (s < t1 + pd.Timedelta(days=1))


# ---------------------------------------------------------------------------
# 固定数据
# ---------------------------------------------------------------------------
def load_koujing(path):
    """基本口径.xls -> (在施工程量 dict, 机构->区 dict)"""
    sheets = pd.read_excel(path, sheet_name=None, header=None)
    zaishi, org2dist = {}, {}
    for name, df in sheets.items():
        head = df.iloc[0].astype(str).tolist()
        if "在施工程量" in head and "市区" in head:
            sub = df.iloc[1:, [1, 2]].dropna()
            for _, r in sub.iterrows():
                try:
                    zaishi[str(r.iloc[0]).strip()] = int(r.iloc[1])
                except (TypeError, ValueError):
                    continue
        if "执法机构" in head and "行政区" in head:
            df2 = df.copy()
            df2.columns = df2.iloc[0]
            df2 = df2.iloc[1:]
            sub = df2[["执法机构", "行政区"]].dropna(subset=["执法机构"])
            org2dist = dict(zip(sub["执法机构"].astype(str).str.strip(),
                                sub["行政区"].astype(str).str.strip()))
    if not zaishi:
        raise ValueError("基本口径.xls 中未找到 在施工程量 表")
    return zaishi, org2dist


def weekly_stats_from_raw(check_raw_path, d0, d1):
    """检查原始导出 -> {区名: dict(check, problem, zgl)}（17区+执法处，按添加时间）。

    口径与 百日攻坚统计表生成程序(deploy) 一致，已对 7.20-7.26 / 7.27-8.2 两周
    逐行验证与 百日攻坚xlsx 完全吻合：
    - 建筑市场检查单：检查单名称（模板） ∈ BM_TEMPLATES
    - 发现问题：处理结果 含 其他/责令改正/一般行政处罚/简易行政处罚
    - 责改立案：处理结果 含 责令改正/一般行政处罚/简易行政处罚
    - 区名：机构名称子串映射，未匹配的归 执法处（市级）
    """
    raw = read_any_excel(check_raw_path)
    sub = raw[day_mask(raw["添加时间"], d0, d1)]
    bm = sub[sub["检查单名称（模板）"].isin(BM_TEMPLATES)]
    dist = bm["机构"].map(extract_district_sub)
    prob = contains_any(bm["处理结果"], RESULT_PROBLEM)
    zgl = contains_any(bm["处理结果"], RESULT_ZGL)
    rows = {}
    for name in DISTRICTS_ORDER + ["执法处"]:
        sel = dist == name
        rows[name] = {"check": int(sel.sum()),
                      "problem": int((sel & prob).sum()),
                      "zgl": int((sel & zgl).sum())}
    return rows


def last_year_check(path, m0, d0, m1, d1):
    """2025年建筑市场检查.xls -> (检查数, 问题数)，按添加时间（与本周口径一致）。"""
    df = pd.read_excel(path, sheet_name=0)
    col = "添加时间" if "添加时间" in df.columns else "检查开始日期"
    mask = md_mask(df[col], m0, d0, m1, d1)
    sub = df[mask]
    n_check = len(sub)
    n_prob = sub["处理结果"].astype(str).str.strip().isin(RESULT_PROBLEM).sum()
    return int(n_check), int(n_prob)


def last_year_punish(path, m0, d0, m1, d1):
    """2025年工程建设领域处罚.xls -> (件数, 罚款金额元)，按添加日期 + 文书类型口径（与本周口径一致）。"""
    df = pd.read_excel(path, sheet_name=0)
    if "文书类型" in df.columns:
        df = df[df["文书类型"].astype(str).str.strip().isin(PUNISH_DOCS)]
    col = "添加日期" if "添加日期" in df.columns else "决定日期"
    mask = md_mask(df[col], m0, d0, m1, d1)
    amt = None
    if "处罚金额" in df.columns:
        amt = float(pd.to_numeric(df.loc[mask, "处罚金额"], errors="coerce").fillna(0).sum())
    return int(mask.sum()), amt


# ---------------------------------------------------------------------------
# 每周原始导出
# ---------------------------------------------------------------------------
def lian_stats(check_raw_path, d0, d1, org2dist):
    """检查原始导出 -> (市级立案数, {top区: 立案数})，按添加时间。"""
    raw = read_any_excel(check_raw_path)
    raw = raw[raw["任务名称"].astype(str).str.contains("建筑市场", na=False)]
    raw["_dist"] = raw["机构"].astype(str).str.strip().map(org2dist)
    mask = day_mask(raw["添加时间"], d0, d1)
    sub = raw[mask]
    is_lian = sub["处理结果"].astype(str).str.strip().isin(RESULT_LIAN)
    lian = sub[is_lian]
    per_dist = lian["_dist"].value_counts().to_dict()
    city = int(per_dist.get("市级", 0))
    return city, {k: int(v) for k, v in per_dist.items()}


def punish_stats(punish_raw_path, d0, d1, org2dist):
    """处罚原始导出 -> (件数, {区: 件数}, 罚款金额元)，按添加日期。"""
    raw = read_any_excel(punish_raw_path)
    raw = raw[raw["文书类型"].astype(str).str.strip().isin(PUNISH_DOCS)]
    raw["_dist"] = raw["执法机构"].astype(str).str.strip().map(org2dist)
    mask = day_mask(raw["添加日期"], d0, d1)
    sub = raw[mask]
    per_dist = sub["_dist"].value_counts().to_dict()
    amt = None
    if "处罚金额" in raw.columns:
        amt = float(pd.to_numeric(sub["处罚金额"], errors="coerce").fillna(0).sum())
    return int(len(sub)), {k: int(v) for k, v in per_dist.items()}, amt


def week_from_exports(*paths):
    """由原始导出中的最大日期推断本周区间（周一~周日），返回 (d0, d1) 或 None。"""
    latest = None
    for p in paths:
        if not p or not os.path.exists(p):
            continue
        try:
            raw = read_any_excel(p)
            for col in ("添加日期", "添加时间"):
                if col in raw.columns:
                    s = pd.to_datetime(raw[col], errors="coerce").dropna()
                    if len(s) and (latest is None or s.max() > latest):
                        latest = s.max()
        except Exception:
            continue
    if latest is None:
        return None
    d = latest.date()
    d1 = d - dt.timedelta(days=(d.weekday() + 1) % 7)   # 最近一个周日
    return d1 - dt.timedelta(days=6), d1


# ---------------------------------------------------------------------------
# 文案生成
# ---------------------------------------------------------------------------
def fmt_pct(cur, prev):
    """环比/同比百分比：四舍五入取整，带方向。返回 ('上升', '20')。"""
    if prev == 0:
        return ("上升", "0")
    v = (cur - prev) / prev * 100
    return ("上升" if v >= 0 else "下降", str(int(round(abs(v)))))


def fmt_pp(rate_cur, rate_prev):
    """百分点差：绝对值<10 保留1位小数，否则取整。"""
    diff = (rate_cur - rate_prev) * 100
    direction = "增长" if diff >= 0 else "下降"
    a = abs(diff)
    if a < 10:
        s = ("%.1f" % a).rstrip("0").rstrip(".")
    else:
        s = str(int(round(a)))
    return direction, s


def fmt_times(cur_rate, last_rate):
    """倍数：1位小数（去尾零）。"""
    if last_rate == 0:
        return "-"
    s = "%.1f" % (cur_rate / last_rate)
    return s.rstrip("0").rstrip(".") if "." in s else s


def short(name):
    return DIST_SHORT.get(name, name)


def fmt_amount(yuan):
    """元 -> 万元字符串（保留1位小数）；None -> '-'。"""
    if yuan is None:
        return "-"
    return "%.1f" % (yuan / 10000)


def build_paragraphs(c):
    """根据计算结果字典 c 生成两段正文。"""
    rng = "%d月%d日-%d月%d日" % (c["m0"], c["d0"], c["m1"], c["d1"])

    hb_dir, hb_val = fmt_pct(c["tw_check"], c["lw_check"])
    tb_dir, tb_val = fmt_pct(c["tw_check"], c["ly_check"])
    tw_rate = c["tw_prob"] / c["tw_check"] if c["tw_check"] else 0
    lw_rate = c["lw_prob"] / c["lw_check"] if c["lw_check"] else 0
    ly_rate = c["ly_prob"] / c["ly_check"] if c["ly_check"] else 0
    pp_dir, pp_val = fmt_pp(tw_rate, lw_rate)
    rate_times = fmt_times(tw_rate, ly_rate)
    city_rate = "%.1f" % (c["city_prob"] / c["city_check"] * 100) if c["city_check"] else "0"
    city_lian = "%.1f" % (c["city_lian"] / c["city_check"] * 100) if c["city_check"] else "0"

    most = short(c["most_dist"])
    seg = "本周在施项目数量较多的前八个区中，%s检查较多" % most
    if c["none_dists"]:
        seg += "，%s未开展检查" % "、".join(short(x) for x in c["none_dists"])
    if c["lian_dists"]:
        names = "、".join(short(x) for x in c["lian_dists"])
        jun = "均" if len(c["lian_dists"]) >= 4 else ""
        seg += "，其中%s%s检查发现问题并拟立案处理" % (names, jun)
    seg += "。"

    para1 = (
        "本周（%s）全市住建系统共开展建筑市场执法检查%d项次（上周%d项次，去年同期项%d次），"
        "检查量环比%s%s%%，同比%s%s%%，"
        "检查发现问题项目%d项次（上周%d项次，去年同期%d项次），"
        "检查发现问题率环比%s%s个百分点，为去年同期的%s倍。"
        "本周市级部门对区级监管项目开展百日攻坚专项检查%d项次，"
        "检查发现问题率%s%%，立案率%s%%。%s"
        % (rng, c["tw_check"], c["lw_check"], c["ly_check"],
           hb_dir, hb_val, tb_dir, tb_val,
           c["tw_prob"], c["lw_prob"], c["ly_prob"],
           pp_dir, pp_val, rate_times,
           c["city_check"], city_rate, city_lian, seg)
    )

    phb_dir, phb_val = fmt_pct(c["tw_punish"], c["lw_punish"])
    ptimes = fmt_times(c["tw_punish"], c["ly_punish"])
    if (c["tw_amount"] is not None and c["lw_amount"] is not None
            and c["ly_amount"] is not None):
        amt_hb_dir, amt_hb_val = fmt_pct(c["tw_amount"], c["lw_amount"])
        amt_times = fmt_times(c["tw_amount"], c["ly_amount"])
        amt_seg = "。罚款金额%s万元，（上周%s万元，去年同期%s万元），环比%s%s%%，为去年同期的%s倍" % (
            fmt_amount(c["tw_amount"]), fmt_amount(c["lw_amount"]),
            fmt_amount(c["ly_amount"]), amt_hb_dir, amt_hb_val, amt_times)
    else:
        amt_seg = ""
    rank_str = "、".join(short(x) for x in c["punish_rank"])
    para2 = (
        "本周（%s）全市住建系统共作出工程建设领域处罚%d件，（上周%d件，去年同期%d件），"
        "环比%s%s%%，为去年同期的%s倍%s。"
        "在施项目数量较多的前八个区中，每百工程处罚量从高到底排名分别是%s。"
        % (rng, c["tw_punish"], c["lw_punish"], c["ly_punish"],
           phb_dir, phb_val, ptimes, amt_seg, rank_str)
    )
    return para1, para2


# ---------------------------------------------------------------------------
# docx 生成（格式与样例一致）
# ---------------------------------------------------------------------------
def write_docx(out_path, para1, para2):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt

    doc = docx.Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Emu(7560310), Emu(10692130)   # A4
    sec.left_margin = sec.right_margin = Emu(1143000)
    sec.top_margin = sec.bottom_margin = Emu(914400)

    def add(text, font_name, size_pt, center=False, indent=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(28)          # 固定行距 28 磅
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if indent:
            ind = p._element.get_or_add_pPr().get_or_add_ind()
            ind.set(qn("w:firstLineChars"), "200")
            ind.set(qn("w:firstLine"), str(int(size_pt * 2 * 20)))
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(size_pt)
        rPr = r._element.get_or_add_rPr()
        rPr.rFonts.set(qn("w:eastAsia"), font_name)
        rPr.rFonts.set(qn("w:cs"), font_name)
        return p

    add("", "方正小标宋简体", 22, center=True)
    add("执法处建筑市场检查周报材料", "方正小标宋简体", 22, center=True)
    add("", "方正小标宋简体", 22, center=True)
    add("本周全市住建系统建筑市场执法情况", "黑体", 16, indent=True)
    add(para1, "仿宋_GB2312", 16, indent=True)
    add("本周全市住建系统工程建设领域处罚情况", "黑体", 16, indent=True)
    add(para2, "仿宋_GB2312", 16, indent=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# 历史记录
# ---------------------------------------------------------------------------
def load_history():
    try:
        with open(HISTORY_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_history(hist):
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(HISTORY_PATH, "w", encoding="utf-8") as f:
        json.dump(hist, f, ensure_ascii=False, indent=1)


def week_key(d0, d1):
    return "%s~%s" % (d0, d1)


# ---------------------------------------------------------------------------
# 主计算流程
# ---------------------------------------------------------------------------
def compute(opts, log):
    """opts: 参数字典；log: 日志函数。返回 (para1, para2)。"""
    d0, d1 = opts["d0"], opts["d1"]                      # 本周起止（date/datetime）
    if isinstance(d0, dt.datetime):
        d0 = d0.date()
    if isinstance(d1, dt.datetime):
        d1 = d1.date()
    lw_d1 = d0 - dt.timedelta(days=1)                    # 上周日
    lw_d0 = lw_d1 - dt.timedelta(days=6)                 # 上周一
    log("本周区间: %s ~ %s；上周区间: %s ~ %s" % (d0, d1, lw_d0, lw_d1))

    zaishi, org2dist = load_koujing(opts["koujing"])
    top8 = sorted(zaishi, key=lambda k: -zaishi[k])[:8]
    log("在施工程量前八区: %s" % "、".join("%s(%d)" % (k, zaishi[k]) for k in top8))

    hist = load_history()
    lw_hist = hist.get(week_key(lw_d0, lw_d1), {})

    tw = weekly_stats_from_raw(opts["check_raw"], d0, d1)
    lw = weekly_stats_from_raw(opts["check_raw"], lw_d0, lw_d1)
    tw_check = sum(v["check"] for v in tw.values())
    tw_prob = sum(v["problem"] for v in tw.values())
    lw_check = sum(v["check"] for v in lw.values())
    lw_prob = sum(v["problem"] for v in lw.values())
    log("本周检查 %d 项次 / 问题 %d 项次；上周 %d / %d（检查导出按添加时间）"
        % (tw_check, tw_prob, lw_check, lw_prob))
    if "check_total" in lw_hist:
        lw_check = lw_hist["check_total"]
        lw_prob = lw_hist.get("problem_total", lw_prob)
        log("上周检查/问题 改采历史记录（与上周报告一致）: %d / %d" % (lw_check, lw_prob))

    ly_check, ly_prob = last_year_check(opts["check2025"], d0.month, d0.day, d1.month, d1.day)
    ly_punish, ly_amount = last_year_punish(opts["punish2025"], d0.month, d0.day, d1.month, d1.day)
    log("去年同期（2025年xls按同月日、添加日期/时间口径）: 检查 %d 次 / 问题 %d 项次 / 处罚 %d 件"
        % (ly_check, ly_prob, ly_punish))
    if ly_amount is not None:
        log("去年同期罚款金额: %.1f 万元" % (ly_amount / 10000))

    city_check = tw["执法处"]["check"]
    city_prob = tw["执法处"]["problem"]

    # 立案数（检查原始导出）
    city_lian, dist_lian = lian_stats(opts["check_raw"], d0, d1, org2dist)
    log("市级专项检查 %d 项次 / 问题 %d；市级立案 %d（检查导出按添加时间）"
        % (city_check, city_prob, city_lian))

    # 前八区分析
    top8_rows = [(name, tw.get(name)) for name in top8]
    most_dist = max(top8_rows, key=lambda kv: (kv[1]["check"] if kv[1] else 0))[0]
    none_dists = [name for name, v in top8_rows if v and v["check"] == 0]
    lian_dists = [name for name, v in top8_rows
                  if dist_lian.get(name, 0) > 0]
    lian_dists.sort(key=lambda n: (-dist_lian.get(n, 0),
                                   list(tw.keys()).index(n)))
    log("检查较多: %s；未开展检查: %s；拟立案区: %s"
        % (most_dist, "、".join(none_dists) or "无",
           "、".join("%s(%d)" % (n, dist_lian.get(n, 0)) for n in lian_dists) or "无"))

    # 处罚
    tw_punish, dist_punish, tw_amount = punish_stats(opts["punish_raw"], d0, d1, org2dist)
    log("本周处罚 %d 件（处罚导出按添加日期）" % tw_punish)
    if tw_amount is not None:
        log("本周罚款金额: %.1f 万元" % (tw_amount / 10000))

    lw_amount = None
    lw_punish = opts.get("lw_punish")
    if lw_punish is not None:
        log("上周处罚 %d 件（手动指定）" % lw_punish)
    elif "punish_total" in lw_hist:
        lw_punish = lw_hist["punish_total"]
        log("上周处罚 %d 件（历史记录，与上周报告一致）" % lw_punish)
    else:
        lw_punish, _, _ = punish_stats(opts["punish_raw"], lw_d0, lw_d1, org2dist)
        log("上周处罚 %d 件（由处罚导出重算，可能与上周报告略有出入）" % lw_punish)
    if lw_punish is not None:
        if "punish_amount" in lw_hist:
            lw_amount = lw_hist["punish_amount"]
            log("上周罚款金额: %.1f 万元（历史记录）" % (lw_amount / 10000))
        elif tw_amount is not None:
            lw_amount = punish_stats(opts["punish_raw"], lw_d0, lw_d1, org2dist)[2]
            if lw_amount is not None:
                log("上周罚款金额: %.1f 万元（由处罚导出重算）" % (lw_amount / 10000))

    punish_rank = sorted(top8, key=lambda n: (-(dist_punish.get(n, 0) / zaishi[n]), -zaishi[n]))
    log("每百工程处罚量: %s" % "、".join(
        "%s(%.2f)" % (n, dist_punish.get(n, 0) / zaishi[n] * 100) for n in punish_rank))

    c = dict(m0=d0.month, d0=d0.day, m1=d1.month, d1=d1.day,
             tw_check=tw_check, lw_check=lw_check, ly_check=ly_check,
             tw_prob=tw_prob, lw_prob=lw_prob, ly_prob=ly_prob,
             city_check=city_check, city_prob=city_prob, city_lian=city_lian,
             most_dist=most_dist, none_dists=none_dists, lian_dists=lian_dists,
             tw_punish=tw_punish, lw_punish=lw_punish, ly_punish=ly_punish,
             tw_amount=tw_amount, lw_amount=lw_amount, ly_amount=ly_amount,
             punish_rank=punish_rank)
    para1, para2 = build_paragraphs(c)

    # 保存历史（供下周环比引用）
    hist[week_key(d0, d1)] = {"punish_total": tw_punish, "check_total": tw_check,
                              "problem_total": tw_prob, "punish_amount": tw_amount,
                              "generated": dt.datetime.now().isoformat()}
    save_history(hist)
    return para1, para2


# ---------------------------------------------------------------------------
# 供 GUI / 网页版共用的入口
# ---------------------------------------------------------------------------
def autodetect_params():
    """自动探测输入文件与参数，返回 dict（尽量填充，可能缺项）。"""
    res = {"check_raw": "", "punish_raw": "", "d0": "", "d1": "",
           "lw_punish": "", "out": ""}

    def newest(patterns):
        files = []
        for pat in patterns:
            for d in (SCRIPT_DIR, os.path.dirname(SCRIPT_DIR)):
                files += [p for p in glob.glob(os.path.join(d, pat))
                          if not os.path.basename(p).startswith("~$")]
        return max(files, key=os.path.getmtime) if files else ""

    res["check_raw"] = newest(["检查数据*.xls", "*检查数据导出*.xls"])
    res["punish_raw"] = newest(["处罚数据*.xls", "*处罚数据导出*.xls"])
    try:
        wk = week_from_exports(res["punish_raw"], res["check_raw"])
        if not wk:
            return res
        d0, d1 = wk
        res["d0"] = d0.isoformat()
        res["d1"] = d1.isoformat()
        res["out"] = os.path.join(SCRIPT_DIR, "执法处周报材料(%d月%d日-%d月%d日).docx"
                                  % (d0.month, d0.day, d1.month, d1.day))
        lw1 = d0 - dt.timedelta(days=1)
        lw0 = lw1 - dt.timedelta(days=6)
        hv = load_history().get(week_key(lw0, lw1), {}).get("punish_total")
        res["lw_punish"] = str(hv) if hv is not None else ""
    except Exception:
        pass
    return res


def run_report(params, log):
    """按参数生成周报 docx，返回 (para1, para2, out_path)。"""
    get = lambda k: str(params.get(k, "")).strip()
    opts = dict(
        check_raw=get("check_raw"), punish_raw=get("punish_raw"),
        check2025=get("check2025") or FIXED_CHECK_XLS,
        punish2025=get("punish2025") or FIXED_PUNISH_XLS,
        koujing=get("koujing") or FIXED_KOUJING_XLS,
        d0=dt.date.fromisoformat(get("d0")), d1=dt.date.fromisoformat(get("d1")),
    )
    for k in ("check_raw", "punish_raw"):
        if not os.path.exists(opts[k]):
            raise ValueError("文件不存在: %s" % opts[k])
    for k in ("check2025", "punish2025", "koujing"):
        if not os.path.exists(opts[k]):
            raise ValueError("固定数据文件缺失: %s" % opts[k])
    lw = get("lw_punish")
    opts["lw_punish"] = int(lw) if lw else None
    out_path = get("out") or os.path.join(SCRIPT_DIR, "执法处周报材料.docx")

    para1, para2 = compute(opts, log)
    write_docx(out_path, para1, para2)
    return para1, para2, out_path


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------
def make_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext, ttk

    app = tk.Tk()
    app.title("执法处周报材料生成器")
    app.geometry("860x720")

    nb = ttk.Notebook(app)
    nb.pack(fill="both", expand=True, padx=6, pady=6)
    frm = ttk.Frame(nb)
    nb.add(frm, text="文件与参数")

    entries = {}

    def row(r, label, key, default="", browse="file"):
        ttk.Label(frm, text=label, width=22, anchor="e").grid(
            row=r, column=0, sticky="e", padx=4, pady=3)
        e = ttk.Entry(frm, width=78)
        e.grid(row=r, column=1, sticky="we", padx=4, pady=3)
        e.insert(0, default)
        entries[key] = e
        if browse == "file":
            ttk.Button(frm, text="浏览…", width=8,
                       command=lambda: e.replace(0, "end", filedialog.askopenfilename(
                           initialdir=SCRIPT_DIR))).grid(row=r, column=2, padx=2)
        elif browse == "save":
            ttk.Button(frm, text="浏览…", width=8,
                       command=lambda: e.replace(0, "end", filedialog.asksaveasfilename(
                           initialdir=SCRIPT_DIR, defaultextension=".docx"))).grid(row=r, column=2, padx=2)

    def autodetect():
        for k, v in autodetect_params().items():
            if v and k in entries:
                entries[k].replace(0, "end", v)

    r = 0
    ttk.Label(frm, text="—— 每周数据（平台原始导出）——", foreground="#555").grid(row=r, column=1, sticky="w"); r += 1
    row(r, "检查原始数据导出:", "check_raw"); r += 1
    row(r, "处罚原始数据导出:", "punish_raw"); r += 1
    ttk.Label(frm, text="—— 参数 ——", foreground="#555").grid(row=r, column=1, sticky="w"); r += 1
    row(r, "本周开始日期:", "d0", "", browse=None); r += 1
    row(r, "本周结束日期:", "d1", "", browse=None); r += 1
    row(r, "上周处罚件数(可留空):", "lw_punish", "", browse=None); r += 1
    row(r, "输出docx:", "out", "", browse="save"); r += 1

    btns = ttk.Frame(frm)
    btns.grid(row=r, column=1, sticky="w", pady=6)
    ttk.Button(btns, text="自动探测文件", command=autodetect).pack(side="left", padx=4)

    logwin = scrolledtext.ScrolledText(app, height=18, font=("Consolas", 10))
    logwin.pack(fill="both", expand=True, padx=6, pady=(0, 6))

    def log(msg):
        logwin.insert("end", str(msg) + "\n")
        logwin.see("end")
        app.update_idletasks()

    def generate():
        logwin.delete("1.0", "end")
        try:
            params = {k: e.get() for k, e in entries.items()}
            para1, para2, out_path = run_report(params, log)
            log("")
            log("【正文一】")
            log(para1)
            log("")
            log("【正文二】")
            log(para2)
            log("")
            log("已生成: %s" % out_path)
            messagebox.showinfo("完成", "周报已生成：\n%s" % out_path)
        except Exception:
            log(traceback.format_exc())
            messagebox.showerror("出错", "生成失败，详见日志窗口。")

    ttk.Button(btns, text="生成周报", command=generate).pack(side="left", padx=4)

    # Entry.replace helper（tk.Entry 无 replace 方法）
    import tkinter as _tk
    def _replace(self, first, last, text):
        self.delete(0, "end")
        self.insert(0, text)
    _tk.Entry.replace = _replace
    ttk.Entry.replace = _replace

    autodetect()
    app.mainloop()


if __name__ == "__main__":
    make_gui()
